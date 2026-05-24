#!/usr/bin/env python3
"""
Script para convertir TEAM-05-FULL-DOCUMENTATION.md a .docx
Genera el archivo en: C:/Users/saule/OneDrive/Documentos/Programación Web/
"""

import re
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def read_md_file(path):
    with open(path, 'r', encoding='utf-8') as f:
        return f.read()

def add_code_block(doc, code):
    """Add a formatted code block."""
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    # Set background to light gray
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    # Add the code with monospace font
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    # Add shading
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'F5F5F5')
    shading.set(qn('w:val'), 'clear')
    p.paragraph_format.element.get_or_add_pPr().append(shading)

def add_table_row(table, cells_data, bold=False, header=False):
    """Add a row to the table."""
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(9)
        if bold or header:
            run.bold = True
        if header:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            # Add dark header background
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), '2B579A')
            shading.set(qn('w:val'), 'clear')
            cell.paragraphs[0].paragraph_format.element.get_or_add_pPr().append(shading)

def convert_md_to_docx(md_content, output_path):
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Configure heading styles
    for i in range(1, 4):
        heading_style = doc.styles[f'Heading {i}']
        heading_style.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)
        heading_style.font.bold = True
    
    doc.styles['Heading 1'].font.size = Pt(22)
    doc.styles['Heading 2'].font.size = Pt(16)
    doc.styles['Heading 3'].font.size = Pt(13)
    
    lines = md_content.split('\n')
    in_code_block = False
    code_buffer = []
    in_table = False
    in_list = False
    table_buffer = []
    
    i = 0
    while i < len(lines):
        line = lines[i]
        
        # Handle code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                add_code_block(doc, '\n'.join(code_buffer))
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue
        
        # Handle tables
        if line.strip().startswith('|') and line.strip().endswith('|'):
            table_buffer.append(line)
            in_table = True
            i += 1
            continue
        else:
            if in_table and len(table_buffer) >= 3:
                # Process table
                # Parse header
                header_cells = [c.strip() for c in table_buffer[0].strip().split('|') if c.strip()]
                # Check if second line is separator
                if '---' in table_buffer[1]:
                    data_start = 2
                else:
                    data_start = 1
                
                # Create table
                num_cols = len(header_cells)
                table = doc.add_table(rows=1, cols=num_cols)
                table.style = 'Light Grid Accent 1'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # Add header row
                for j, header in enumerate(header_cells):
                    cell = table.rows[0].cells[j]
                    cell.text = ''
                    p = cell.paragraphs[0]
                    run = p.add_run(header)
                    run.bold = True
                    run.font.size = Pt(9)
                
                # Add data rows
                for row_idx in range(data_start, len(table_buffer)):
                    row_cells = [c.strip() for c in table_buffer[row_idx].strip().split('|') if c.strip()]
                    if len(row_cells) != num_cols:
                        continue
                    row = table.add_row()
                    for j, cell_text in enumerate(row_cells):
                        cell = row.cells[j]
                        cell.text = ''
                        p = cell.paragraphs[0]
                        run = p.add_run(cell_text)
                        run.font.size = Pt(9)
                
                doc.add_paragraph()  # spacing
                table_buffer = []
                in_table = False
            elif in_table:
                table_buffer = []
                in_table = False
        
        # Handle headings
        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:].strip(), level=4)
        elif line.startswith('- ') or line.startswith('* '):
            # Handle list items
            text = line[2:].strip()
            # Check for bold markers **...**
            if '**' in text:
                p = doc.add_paragraph(style='List Bullet')
                parts = re.split(r'(\*\*.*?\*\*)', text)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    else:
                        p.add_run(part)
            else:
                doc.add_paragraph(text, style='List Bullet')
        elif line.strip() == '---':
            # Horizontal rule
            p = doc.add_paragraph()
            run = p.add_run('_' * 60)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            run.font.size = Pt(8)
        elif line.strip() == '':
            if not in_list:
                pass  # Skip empty lines
            in_list = False
        else:
            # Normal paragraph text
            text = line.strip()
            if text:
                # Check for bold markers
                if '**' in text:
                    p = doc.add_paragraph()
                    parts = re.split(r'(\*\*.*?\*\*)', text)
                    for part in parts:
                        if part.startswith('**') and part.endswith('**'):
                            run = p.add_run(part[2:-2])
                            run.bold = True
                        else:
                            p.add_run(part)
                else:
                    p = doc.add_paragraph(text)
                in_list = False
        
        i += 1
    
    # Handle any remaining code block
    if in_code_block and code_buffer:
        add_code_block(doc, '\n'.join(code_buffer))
    
    # Save
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"✅ Documento DOCX generado en: {output_path}")
    print(f"   Tamaño: {os.path.getsize(output_path) / 1024:.1f} KB")

if __name__ == '__main__':
    md_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'docs', 'TEAM-05-FULL-DOCUMENTATION.md')
    output_path = r'C:\Users\saule\OneDrive\Documentos\Programación Web\TEAM-05-TurboPapus-Documentacion-Completa.docx'
    
    if not os.path.exists(md_path):
        print(f"❌ No se encontró el archivo .md: {md_path}")
        exit(1)
    
    md_content = read_md_file(md_path)
    convert_md_to_docx(md_content, output_path)
